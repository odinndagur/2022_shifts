#!/Users/odinndagur/.odb/vaktaplan/camelot-env/bin/python3


# IMPORTS
#region [blue2]
import camelot
import pandas as pd
import os
import sys
import re
from numpy import nan
from pathlib import Path
import pdfplumber
import datetime
from collections import defaultdict
import docx
from math import ceil,floor
from pdf2image import convert_from_path

import argparse
import sys
#endregion

class Vaktaplan():
    def __init__(self,df:pd.DataFrame):
        self.df = df
        self.people = self.get_people()
        self.dayplans = self.generate_dayplans()
        self.start_date, self.end_date = self.get_date_range()

    def __str__(self):
        ppl_list = '\n'.join(self.people)
        return f'Shiftplan from {self.start_date} to {self.end_date}.\nPeople:\n{ppl_list}'

    def generate_dayplans(self):
        days = list(self.get_days())
        for idx, day in enumerate(days):
            temp_date = day['date']['date'].split('\n')[0]
            print(f'Making dayplan for {temp_date}')
            yield self.doc_from_date_day(day)
    
    def get_days(self): #yields generator
        for idx in range(len(self.df.columns)):
            day = defaultdict(lambda: defaultdict(list))
            slice = self.df.iloc[:,idx].replace('',nan).dropna()
            if not len(slice):
                continue
            day['date']['date'] = slice.name
            # date = slice.name
            for row_idx, shift in enumerate(slice):
                person = self.get_name(slice.index[row_idx])
                if all([shift != 'ORLOF',shift]):
                    time, type = shift.split(' ')
                    if type == 'UB':
                        day['UB']['ub'].append(' '.join([person,time]))
                    else:
                        starttime,endtime = time.split('-')
                        starthr = int(starttime.split(':')[0])
                        endhr = int(endtime.split(':')[0])
                        if 0 < starthr < 12:
                            if endhr > 16:
                                col = 'dv'
                            else:
                                col = 'mv'
                        if 12 <= starthr < 16:
                            col = 'dv'
                        if 16 <= starthr < 23:
                            col = 'kv'
                        if 23 <= starthr <= 24:
                            col = 'nv'
                        if 20 < endhr <= 24:
                            col = 'kv'
                        # print(person,starthr,endhr,type,col)
                        day[type][col].append(','.join([person,time,type,col]))
            yield day

    @classmethod
    def pdf_to_df(cls,file: str) -> pd.DataFrame:
        print('processing pdf')
        global h,w,new_h,new_w,pdfs
        pdfs = convert_from_path(file)
        with pdfplumber.open(file) as pdf:
            page_1 = pdf.pages[0]
        h,w = page_1.height, page_1.width
        new_h,new_w = pdfs[0].height, pdfs[0].width

        tables = camelot.read_pdf(file,pages='1-end',flavor='lattice',line_scale=50,line_tol=1)
        cls.add_shift_text(tables)
        processed_dfs = [cls.process_df(table) for table in tables]
        # concat fyrst
        concatenated_dfs = [pd.concat(processed_dfs[offset:offset+cls.get_num_pages(tables)]) for offset in range(0,tables.n,cls.get_num_pages(tables))]
        # síðan join
        output_df = concatenated_dfs[0]
        for df in concatenated_dfs[1:]:
            output_df = output_df.join(df)
        return output_df

    @classmethod
    def doc_from_date_day(cls,day) -> docx.Document:
        date = day['date']['date']
        doc = docx.Document(os.path.join(os.path.dirname(os.path.abspath(__file__)),'fim_proto.docx'))
        for col_idx,col in enumerate(doc.tables[1].columns):
            for cell in doc.tables[1].column_cells(col_idx):
                if cell.text.strip() == 'dags':
                    p = cell.paragraphs[0]
                    p.clear()
                    p.add_run(date.split('\n')[0])
                if cell.text.strip() == 'UB':
                    p = cell.paragraphs[0]
                    p.clear()
                    p.add_run('\n'.join(day['UB']['ub']))
                if cell.text.strip() == 'vikudags':
                    p = cell.paragraphs[0]
                    p.clear()
                    decimal_date = date.split('\n')[0]
                    if decimal_date:
                        date_day, date_month = decimal_date.split('.')
                        weekday = cls.get_weekday(month = date_month, day = date_day)
                        p.add_run(weekday)
                if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                    shift = cell.text.split(' ')
                    if not shift[0] in day:
                        cell.paragraphs[0].clear()
                    else:
                        if shift[1] in day[shift[0]]:
                            temp = ' ' if 'NV' in shift[0] else '\n'
                            ppl = [
                                p.split(',')[0]+ temp + p.split(',')[1] for p in day[shift[0]][shift[1]]
                            ]
                            ppl = sorted(ppl, key=lambda x: int(x.split(temp)[-1].split(':')[0]))
                            p = cell.paragraphs[0]
                            p.clear()
                            p.add_run('\n'.join(ppl))


        for col_idx,col in enumerate(doc.tables[1].columns):
            for cell in doc.tables[1].column_cells(col_idx):
                if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                    cell.paragraphs[0].clear()
        date_str = date.split('\n')[0]
        # output_filename = os.path.join(output_directory,f'{date_str} editad.docx')
        # doc.save(output_filename)
        # os.system(f'open {os.path.join(os.getcwd(),output_filename)}')
        return doc
    
    @classmethod
    def get_shifts_for_person(df: pd.DataFrame,person: str) -> str:
        shifts = df.loc[person].replace('',nan).dropna()
        return [f'{date.split(chr(10))[0]} {shift}' for date,shift in zip(shifts.index,shifts.values)]

    def get_people(self) -> list[str]:
        return [p for p in list(self.df.index) if len(p)]

    def get_days(self): #yields generator
        for idx in range(len(self.df.columns)):
            day = defaultdict(lambda: defaultdict(list))
            slice = self.df.iloc[:,idx].replace('',nan).dropna()
            if not len(slice):
                continue
            day['date']['date'] = slice.name
            # date = slice.name
            for row_idx, shift in enumerate(slice):
                person = self.get_name(slice.index[row_idx])
                if all([shift != 'ORLOF',shift]):
                    time, type = shift.split(' ')
                    if type == 'UB':
                        day['UB']['ub'].append(' '.join([person,time]))
                    else:
                        starttime,endtime = time.split('-')
                        starthr = int(starttime.split(':')[0])
                        endhr = int(endtime.split(':')[0])
                        if 0 < starthr < 12:
                            if endhr > 16:
                                col = 'dv'
                            else:
                                col = 'mv'
                        if 12 <= starthr < 16:
                            col = 'dv'
                        if 16 <= starthr < 23:
                            col = 'kv'
                        if 23 <= starthr <= 24:
                            col = 'nv'
                        if 20 < endhr <= 24:
                            col = 'kv'
                        # print(person,starthr,endhr,type,col)
                        day[type][col].append(','.join([person,time,type,col]))
            yield day

    @classmethod
    def get_weekday(cls,year:int = None, month:int = None, day:int = None) -> str:
        days = ["Mánudagur", "Þriðjudagur", "Miðvikudagur", "Fimmtudagur", "Föstudagur", "Laugardagur", "Sunnudagur"]
        if not any([year,month,day]):
            raise ValueError('Need to input date')
        currentDate = datetime.datetime.today()
        if not year:
            if int(month) < currentDate.month and currentDate.month > 10:
                year = currentDate.year + 1
            else:
                year = currentDate.year
        inputDate = datetime.date(year=year,month=int(month),day=int(day))
        return days[inputDate.weekday()]

    @classmethod
    def is_first_page(df: pd.DataFrame) -> bool:
        for x in df.iloc[:,0].values:
            if 'Hæf' in x:
                return True
        return False

    def get_first_date_cell(self) -> tuple[int, int]:
        import re
        for row_idx,row in self.df.iterrows():
            for col_idx, cell in enumerate(row):
                # print(f'row {row_idx}, col {col_idx}, cell: {cell}')
                if re.match('[0-9][0-9]\.[0-9][0-9]', cell):
                    return col_idx,row_idx

    @classmethod
    def get_num_pages(cls,tables):
        counts = []
        last_first_page = 0
        for idx,table in enumerate(tables):
            if cls.is_first_page(table.df):
                if idx > 0:
                    counts.append(idx - last_first_page)
                    last_first_page = idx
        if len(set(counts)) == 1:
            return counts[0]
        else:
            return counts

    @classmethod
    def get_color(cls,img,cell):
        y = new_h - ((cell.y1 + 3)/h * new_h)
        x = ((cell.x1 + cell.x2)/2)/w * new_w
        return img.getpixel((x,y))

    @classmethod
    def get_colors_from_tables(cls,tables):
        colors = set()
        for idx, table in enumerate(tables):
            for row in table.cells:
                for cell in row:
                    colors.add(cls.get_color(pdfs[idx],cell))
        return colors

    colors =  {
        (255, 255, 0): 'GH', #gulur
        (198, 198, 198): '',#ljosgrar
        (240, 240, 240): '',#ljosljosgrar
        (255, 128, 255): 'NV',
        (0, 128, 0): 'BEG',#graenn
        (255, 0, 0): 'GR',#raudur
        (129, 129, 129): '',#mediumgrar
        (122, 122, 122): '',#mediumgrar
        (128, 0, 255): 'LRL',#fjolublar
        (255, 255, 255): 'UB',#hvitur
        (80, 138, 160): 'ORLOF',
        (128, 128, 64): '',
        (128, 128, 128): 'PHA',#mediumgrar
        (128, 0, 64): 'AS',
        }

    @classmethod
    def add_shift_text(cls,tables) -> None:
        for page_num,table in enumerate(tables):
            df = table.df
            for row_num,row in enumerate(table.cells):
                for col_num,cell in enumerate(row):
                    cell_color = cls.get_color(pdfs[page_num],cell=cell)
                    if cell_color in cls.colors and (cell.text != '') and re.match('[0-9][0-9]:[0-9][0-9]-[0-9][0-9]:[0-9][0-9]',cell.text):
                        df.iloc[row_num,col_num] += ' ' + cls.colors[cell_color]

    @classmethod
    def process_df(cls,table) -> pd.DataFrame:
        df = table.df.copy()
        x,y = cls.get_first_date_cell(df)
        df.iloc[y,x-1] = 'Starfsmaður'
        df.columns = df.iloc[y,:]
        df = df.iloc[y+1:,x-1:]
        df = df.set_index('Starfsmaður')
        df = df.replace('',nan).dropna(how='all').dropna(how='all',axis=1).replace(nan,'')
        return df

    @classmethod     
    def get_name(cls,person: str):
        from ppl import ppl
        if person in ppl:
            if ppl[person]:
                return ppl[person]
        return person

    @staticmethod
    def from_pdf(file:str):
        plan = Vaktaplan(df=Vaktaplan.pdf_to_df(file))
        return plan

    @staticmethod
    def from_csv(file:str):
        plan = Vaktaplan(df = pd.read_csv(file,index_col=0,header=0))
        return plan

    @staticmethod
    def from_pickle(file:str):
        pass

    def get_date_range(self):
        first_day, first_month = self.df.columns[0].split('\n')[0].split('.')
        last_day, last_month = self.df.columns[-1].split('\n')[0].split('.')
        currentDate = datetime.datetime.today()
        if int(first_month) < currentDate.month and currentDate.month > 10:
            year = currentDate.year + 1
        else:
            year = currentDate.year
        first_date = datetime.date(year=year,month=int(first_month),day=int(first_day))
        if int(last_month) < currentDate.month and currentDate.month > 10:
            year = currentDate.year + 1
        else:
            year = currentDate.year
        last_date = datetime.date(year=year,month=int(last_month),day=int(last_day))
        return first_date,last_date
