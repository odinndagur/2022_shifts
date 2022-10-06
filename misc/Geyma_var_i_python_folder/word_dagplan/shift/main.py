from docx import Document
import pandas as pd
import numpy as np

def getdays(df):
    days = []
    for colindex in range(len(df.columns)):
        # name = df.iloc[colindex,1]
        date = ''
        day = ''
        if df.columns[colindex][0].isdigit(): #ef er dagur en ekki starfsmaður eða tomt
            temp = df.columns[colindex] #finna dagsetningu
            if temp[0].isdigit():
                # print(temp)
                temp = temp.split('\n')
                date = temp[0] #11.10 td
                day = temp[1] #mið t.d.
                cols = df.iloc[:,colindex] #allir cells í þessum degi
                shifts = {} #dict til að halda utanum hver er a hvaða vakt
                for j in range(len(cols)): #iteratea yfir vaktir dagsins
                    if isinstance(cols[j],str): #ef er vakt en ekki nan
                        col = cols[j]
                        # print(date, col)
                        name = df.iloc[j,1] #nafn a starfsmanni
                        shift = col.split('\n')[0] #timasetning vaktar
                        shifttype = col.split('\n')[1] #tegund vaktar
                        # print(date,day,name,shift, shifttype)
                        shifts[name] = { #object fyrir stm a þennan dag
                            'shift':shift,
                            'shifttype':shifttype,
                        }
                        # print(name, shifts[name])
                days.append({
                    'date':date,
                    'day': day,
                    'shifts':shifts
                })
    return days

def getshorttime(sh):
    shorttime = ''
    if sh[0].isdigit():
        if sh[0] == '0':
            sh = sh[1:]
        sh = sh.split('-')
        # print(sh)
        start = sh[0].split(':')
        end = sh[1].split(':')
        # print(start,end)
        starthr = start[0]
        endhr = end[0]
        shorttime = starthr + '-' + endhr
        # print(hr,mins)    
        # print(hrs)
    return shorttime

def getcellofshift(shift,t):
    if shift['shifttype'] in names.keys():
        inputname = names[shift['shifttype']]
    inputshift = shift['shift']
    shortshift = getshorttime(inputshift)

    output = {'table':'','cell':''}

    for cidx in range(len(t.columns)):
        col = table.columns[cidx]
        name = col.cells[2].text
        # print(name)
        for cellidx in range(len(col.cells)):
            cell = col.cells[cellidx]
            text = cell.text.strip()
            if cellidx < 9:
                name = table.columns[cidx].cells[2].text
                # print(name)
            if cellidx == 9:
                name = 'NV'
                # print(name)
            if text == inputshift and name == inputname:
                output['table'] = cidx
                output['cell'] = cellidx
    return output

names = {
    'LRL':'Linda',
    'PHA':'Páll',
    'JS':'Jökull',
    'GR':'Garðar',
    'BÓB':'Bragi',
    'GH':'Gunndís',
    'NV':'Næturvakt',
    'ORLOF':'Orlof',
    'Undirbúningur':0,
}

df = pd.read_csv('ALLT.csv')
days = getdays(df)
# print(days)

day = days[0]
shifts = day['shifts']
# print(shifts)

daytuples = []
for i in range(7):
    day = days[i]
    daytuples[i] = {}
    doc = Document(days[i]['day'] + '.docx')
    table = doc.tables[1]
    for cidx in range(len(table.columns)):
        col = table.columns[cidx]
        name = col.cells[2].text
        # print(name)
        for cellidx in range(len(col.cells)):
            cell = col.cells[cellidx]
            text = cell.text.strip()
            if cellidx < 9:
                name = table.columns[cidx].cells[2].text
            if cellidx == 9:
                name = 'NV'
            if len(text) > 1 and text[0].isdigit():
                if not daytuples[i][name]:
                    daytuples[i][name] = []
                daytuples[i][name].append(cell.text,cidx,cellidx)
            #     print(cidx, cellidx, name, text)
print(daytuples)
dayindex = 0
for day in days:
    doc = Document(days[dayindex]['day'] + '.docx')
    table = doc.tables[1]
    for cidx in range(len(table.columns)):
        col = table.columns[cidx]
        name = col.cells[2].text
        # print(name)
        for cellidx in range(len(col.cells)):
            cell = col.cells[cellidx]
            text = cell.text.strip()
            if cellidx < 9:
                name = table.columns[cidx].cells[2].text
            if cellidx == 9:
                name = 'NV'
            # if len(text) > 1 and text[0].isdigit():
            #     print(cidx, cellidx, name, text)

    # print(table.columns[3].cells[2].text)
    # columnindex = 0
    # for col in table.columns:
    #     cellindex = 0
    #     for cell in col.cells:
    #         # print(columnindex, cellindex, cell.text)
    #         if len(cell.text) > 1:
    #             if cell.text[0].isdigit():
    #                 # print('txt: ', cell.text, '  og hitt: ', table.columns[columnindex].cells[cellindex].text)
    #                 print(columnindex, cellindex, cell.text.strip(),table.columns[columnindex].cells[cellindex].text)
    #     # print(col)
    #     columnindex += 1



    # print(getcellofshift(days[0]['shifts']['André Bachmann\n']))
    # print(days[5])
    # print(table.columns[0].cells)
    # for row in table.rows:
    #     for cell in row.cells:
    #         for paragraph in cell.paragraphs:
    #             # print(paragraph.text, cell, row, table)


    #             for person, shift in days[0]['shifts'].items():
    #                 # print(shift['shifttype'])
    #                 # print(counts[shift['shifttype']])
    #                 # print(person, shift)
                    
    #                 # print(shift['shift'])
    #                 shorttime = getshorttime(shift['shift'])
    #                 # print(s,c)
    #                 # print(hrs[0])
    #                 s = shift['shifttype']
    #                 c = str(counts[shift['shifttype']])
    #                 t = s + c + ',' + shorttime
    #                 if t in paragraph.text:
    #                     # print(person, '|' ,t, '|', paragraph.text)
    #                     paragraph.text = paragraph.text.replace(t,person + ' ' + shorttime,1)
    #                     counts[shift['shifttype']] +=1
    #                     # print(person,counts[shift['shifttype']])
    #                     continue
    #             # if 'LRL0,' in paragraph.text:
    #             #     paragraph.text = paragraph.text.replace("LRL0,", "lolz")
    # # print(doc)

        # print(b)
    for a,b in day['shifts'].items():
        name = a
        shift = b['shift']
        shifttype = b['shifttype']
        print(getcellofshift(b,table))
        # incol,incell = getcellofshift(b)
        # table.columns[incol].cells[incell].text = name + shift
        # print(day['date'], a.strip(), b['shift'], b['shifttype'], getshorttime(b['shift']))
    filename = day['date']
    filename = filename.split('.')
    filename = filename[1] + '_' + filename[0] + '.docx'
    doc.save('output/' + filename)