import docx
import pandas as pd
import re
import os
from weekday import get_weekday
from collections import defaultdict
import sys

# doc = docx.Document('fim_proto.docx')


input_filename = sys.argv[1] if sys.argv[1].endswith('.csv') else 'Vaktaplan 11.07-10.08 lokaútgáfa.csv'
df = pd.read_csv(input_filename)
df.index = df[df.columns[0]]
df = df.drop('Starfsmaður',axis=1)


def get_days(df):
    for idx in range(len(df.columns)):
        day = defaultdict(lambda: defaultdict(list))
        slice = df.iloc[:,idx].dropna()
        day['date']['date'] = slice.name
        # date = slice.name
        for row_idx, shift in enumerate(slice):
            person = slice.index[row_idx]
            if shift != 'ORLOF':
                time, type = shift.split(' ')
                if type == 'UB':
                    day['UB']['ub'].append(' '.join([person,time]))
                else:
                    starttime,endtime = time.split('-')
                    starthr = int(starttime.split(':')[0])
                    endhr = int(endtime.split(':')[0])
                    if 0 < starthr < 12:
                        if endhr > 16:
                            col = 'dv'
                        else:
                            col = 'mv'
                    if 12 <= starthr < 16:
                        col = 'dv'
                    if 16 <= starthr < 23:
                        col = 'kv'
                    if 23 <= starthr <= 24:
                        col = 'nv'
                    if 20 < endhr <= 24:
                        col = 'kv'
                    print(person,starthr,endhr,type,col)
                    day[type][col].append(','.join([person,time,type,col]))
        yield day

# day = defaultdict(lambda: defaultdict(list))
# date = df.columns[1]
# for person,shift in df.iloc[:,0:2].dropna().values:
#     if shift != 'ORLOF':
#         time, type = shift.split(' ')
#         if type == 'UB':
#             day['UB']['ub'].append(' '.join([person,time]))
#         else:
#             starttime,endtime = time.split('-')
#             starthr = int(starttime.split(':')[0])
#             endhr = int(endtime.split(':')[0])
#             if 0 < starthr < 12:
#                 if endhr > 16:
#                     col = 'dv'
#                 else:
#                     col = 'mv'
#             if 12 <= starthr < 16:
#                 col = 'dv'
#             if 16 <= starthr < 23:
#                 col = 'kv'
#             if 23 <= starthr <= 24:
#                 col = 'nv'
#             if 20 < endhr <= 24:
#                 col = 'kv'
#             print(person,starthr,endhr,type,col)
#             day[type][col].append(','.join([person,time,type,col]))


# for col_idx,col in enumerate(doc.tables[1].columns):
#     for cell in doc.tables[1].column_cells(col_idx):
#         if not any([cell.text == 'ORLOF', cell.text == '\xa0',len(cell.text.strip()) < 1]):
#             shift = cell.text.split(' ')
#             if shift[0] in day:
#                 # print(cell.text)
#                 if shift[1] in day[shift[0]]:
#                     ppl = [' '.join(p.split(',')[:2]) for p in day[shift[0]][shift[1]]]
#                     ppl = sorted(ppl, key=lambda x: int(x.split(' ')[-1].split(':')[0]))
#                     p = cell.paragraphs[0]
#                     p.clear()
#                     p.add_run('\n'.join(ppl))
#                     # cell.text.
#                     # print(cell.text,'\n\t',day[shift[0]][shift[1]])
#                 else:
#                     if re.match('\D{2,3} \D{2}',cell.text):
#                         cell.paragraphs[0].clear()
def doc_from_date_day(day):
    print(f'day {day}')
    date = day['date']['date']
    doc = docx.Document('fim_proto.docx')
    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if cell.text.strip() == 'dags':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run(date.split('\n')[0])
            if cell.text.strip() == 'UB':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run('\n'.join(day['UB']['ub']))
            if cell.text.strip() == 'vikudags':
                p = cell.paragraphs[0]
                p.clear()
                decimal_date = date.split('\n')[0]
                date_day, date_month = decimal_date.split('.')
                weekday = get_weekday(month = date_month, day = date_day)
                p.add_run(weekday)
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                shift = cell.text.split(' ')
                if not shift[0] in day:
                    cell.paragraphs[0].clear()
                else:
                    # print(cell.text)
                    if shift[1] in day[shift[0]]:
                        # ppl = [' '.join(p.split(',')[:2]) for p in day[shift[0]][shift[1]]]
                        temp = ' ' if 'NV' in shift[0] else '\n'
                        ppl = [
                            p.split(',')[0].split(' ')[0] + temp + p.split(',')[1] for p in day[shift[0]][shift[1]]
                        ]
                        ppl = sorted(ppl, key=lambda x: int(x.split(temp)[-1].split(':')[0]))
                        p = cell.paragraphs[0]
                        p.clear()

                        # for p in ppl:

                        p.add_run('\n'.join(ppl))
                        # cell.text.
                        # print(cell.text,'\n\t',day[shift[0]][shift[1]])
                    # else:
                    #     cell.paragraphs[0].clear()
                # else:
                #     cell.paragraphs[0].clear()

    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                cell.paragraphs[0].clear()
    date_str = date.split('\n')[0]
    output_filename = os.path.join('output',f'{date_str} editad.docx')
    doc.save(output_filename)
    # os.system(f'open {os.path.join(os.getcwd(),output_filename)}')

for x in [(col_idx, row_idx, cell.text) for col_idx, column in enumerate(doc.tables[1].columns) for row_idx, cell in enumerate(column.cells)]:
    print(x)