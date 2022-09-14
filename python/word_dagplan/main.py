import docx
import pandas as pd

import re

doc = docx.Document('manlol.docx')

df = pd.read_csv('/Users/odinndagur/Downloads/vaktaplans_test/Vaktaplan 11.07-10.08 lokaútgáfa.csv')

from collections import defaultdict
day = defaultdict(lambda: defaultdict(list))

for person,shift in df.iloc[:,0:2].dropna().values:
    if shift != 'ORLOF':
        time, type = shift.split(' ')
        starttime,endtime = time.split('-')
        starthr = int(starttime.split(':')[0])
        endhr = int(endtime.split(':')[0])
        if 0 < starthr < 12:
            if endhr > 16:
                col = 'dv'
            else:
                col = 'mv'
        if 12 <= starthr < 16:
            col = 'dv'
        if 16 <= starthr < 23:
            col = 'kv'
        if 23 <= starthr <= 24:
            col = 'nv'
        if 20 < endhr <= 24:
            col = 'kv'
        print(person,starthr,endhr,type,col)
        day[type][col].append(','.join([person,time,type,col]))


# for col_idx,col in enumerate(doc.tables[1].columns):
#     for cell in doc.tables[1].column_cells(col_idx):
#         if not any([cell.text == 'ORLOF', cell.text == '\xa0',len(cell.text.strip()) < 1]):
#             shift = cell.text.split(' ')
#             if shift[0] in day:
#                 # print(cell.text)
#                 if shift[1] in day[shift[0]]:
#                     ppl = [' '.join(p.split(',')[:2]) for p in day[shift[0]][shift[1]]]
#                     ppl = sorted(ppl, key=lambda x: int(x.split(' ')[-1].split(':')[0]))
#                     p = cell.paragraphs[0]
#                     p.clear()
#                     p.add_run('\n'.join(ppl))
#                     # cell.text.
#                     # print(cell.text,'\n\t',day[shift[0]][shift[1]])
#                 else:
#                     if re.match('\D{2,3} \D{2}',cell.text):
#                         cell.paragraphs[0].clear()
            
for col_idx,col in enumerate(doc.tables[1].columns):
    for cell in doc.tables[1].column_cells(col_idx):
        if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
            shift = cell.text.split(' ')
            if not shift[0] in day:
                cell.paragraphs[0].clear()
            else:
                # print(cell.text)
                if shift[1] in day[shift[0]]:
                    ppl = [' '.join(p.split(',')[:2]) for p in day[shift[0]][shift[1]]]
                    ppl = sorted(ppl, key=lambda x: int(x.split(' ')[-1].split(':')[0]))
                    p = cell.paragraphs[0]
                    p.clear()
                    p.add_run('\n'.join(ppl))
                    # cell.text.
                    # print(cell.text,'\n\t',day[shift[0]][shift[1]])
                # else:
                #     cell.paragraphs[0].clear()
            # else:
            #     cell.paragraphs[0].clear()

doc.save('/Users/odinndagur/Desktop/editad.docx')

for x in [(col_idx, row_idx, cell.text) for col_idx, column in enumerate(doc.tables[1].columns) for row_idx, cell in enumerate(column.cells)]:
    print(x)