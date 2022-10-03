import camelot
import pandas as pd
import os
import sys
import re
from numpy import nan
from pathlib import Path
import pdfplumber
import datetime

from collections import defaultdict
import sys
import docx


output_directory = os.path.join(Path(__file__).parent.resolve().parent.resolve(),'output')
file = sys.argv[1] if sys.argv[1].endswith('pdf') else '/Users/odinndagur/Downloads/vaktaplans_test/vaktaplan11.09-10.10 lokaútgáfa.pdf'
filename = os.path.basename(file)
stripped_filename = os.path.splitext(filename)[0]

print(f'running on file: {file}')


def get_days(df): #yields generator
    for idx in range(len(df.columns)):
        day = defaultdict(lambda: defaultdict(list))
        slice = df.iloc[:,idx].dropna()
        day['date']['date'] = slice.name
        # date = slice.name
        for row_idx, shift in enumerate(slice):
            person = slice.index[row_idx]
            if all([shift != 'ORLOF',shift]):
                time, type = shift.split(' ')
                if type == 'UB':
                    day['UB']['ub'].append(' '.join([person,time]))
                else:
                    starttime,endtime = time.split('-')
                    starthr = int(starttime.split(':')[0])
                    endhr = int(endtime.split(':')[0])
                    if 0 < starthr < 12:
                        if endhr > 16:
                            col = 'dv'
                        else:
                            col = 'mv'
                    if 12 <= starthr < 16:
                        col = 'dv'
                    if 16 <= starthr < 23:
                        col = 'kv'
                    if 23 <= starthr <= 24:
                        col = 'nv'
                    if 20 < endhr <= 24:
                        col = 'kv'
                    print(person,starthr,endhr,type,col)
                    day[type][col].append(','.join([person,time,type,col]))
        yield day

def doc_from_date_day(day):
    print(f'day {day}')
    date = day['date']['date']
    doc = docx.Document('python/fim_proto.docx')
    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if cell.text.strip() == 'dags':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run(date.split('\n')[0])
            if cell.text.strip() == 'UB':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run('\n'.join(day['UB']['ub']))
            if cell.text.strip() == 'vikudags':
                p = cell.paragraphs[0]
                p.clear()
                decimal_date = date.split('\n')[0]
                if decimal_date:
                    date_day, date_month = decimal_date.split('.')
                    weekday = get_weekday(month = date_month, day = date_day)
                    p.add_run(weekday)
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                shift = cell.text.split(' ')
                if not shift[0] in day:
                    cell.paragraphs[0].clear()
                else:
                    # print(cell.text)
                    if shift[1] in day[shift[0]]:
                        # ppl = [' '.join(p.split(',')[:2]) for p in day[shift[0]][shift[1]]]
                        temp = ' ' if 'NV' in shift[0] else '\n'
                        ppl = [
                            p.split(',')[0].split(' ')[0] + temp + p.split(',')[1] for p in day[shift[0]][shift[1]]
                        ]
                        ppl = sorted(ppl, key=lambda x: int(x.split(temp)[-1].split(':')[0]))
                        p = cell.paragraphs[0]
                        p.clear()

                        # for p in ppl:

                        p.add_run('\n'.join(ppl))
                        # cell.text.
                        # print(cell.text,'\n\t',day[shift[0]][shift[1]])
                    # else:
                    #     cell.paragraphs[0].clear()
                # else:
                #     cell.paragraphs[0].clear()

    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                cell.paragraphs[0].clear()
    date_str = date.split('\n')[0]
    output_filename = os.path.join('output',f'{date_str} editad.docx')
    doc.save(output_filename)
    # os.system(f'open {os.path.join(os.getcwd(),output_filename)}')
    return doc

def get_weekday(year=None, month=None, day=None):
    days = ["Mánudagur", "Þriðjudagur", "Miðvikudagur", "Fimmtudagur", "Föstudagur", "Laugardagur", "Sunnudagur"]
    if not any([year,month,day]):
        raise ValueError('Need to input date')
    currentDate = datetime.datetime.today()
    if not year:
        if int(month) < currentDate.month and currentDate.month > 10:
            year = currentDate.year + 1
        else:
            year = currentDate.year
    inputDate = datetime.date(year=year,month=int(month),day=int(day))
    return days[inputDate.weekday()]

def is_first_page(df):
    for x in df.iloc[:,0].values:
        if 'Hæf' in x:
            return True
    return False

def get_first_date_cell(df):
    import re
    for row_idx,row in df.iterrows():
        for col_idx, cell in enumerate(row):
            # print(f'row {row_idx}, col {col_idx}, cell: {cell}')
            if re.match('[0-9][0-9]\.[0-9][0-9]', cell):
                return col_idx,row_idx

def get_num_pages(tables):
    counts = []
    last_first_page = 0
    for idx,table in enumerate(tables):
        if is_first_page(table.df):
            if idx > 0:
                counts.append(idx - last_first_page)
                last_first_page = idx
    if len(set(counts)) == 1:
        return counts[0]
    else:
        return counts

def get_pdf_size(filepath):
    with pdfplumber.open(filepath) as pdf:
        page_1 = pdf.pages[0]
    return page_1.height, page_1.width

from pdf2image import convert_from_path
pdfs = convert_from_path(file)

h,w = get_pdf_size(file)
new_h,new_w = pdfs[0].height, pdfs[0].width

def get_color(img,cell):
    y = new_h - ((cell.y1 + 3)/h * new_h)
    x = ((cell.x1 + cell.x2)/2)/w * new_w
    return img.getpixel((x,y))

# tables = camelot.read_pdf(file,pages='1-end',flavor='lattice',line_scale=20,line_tol=5)
tables = camelot.read_pdf(file,pages='1-end',flavor='lattice',line_scale=50,line_tol=1)

def get_colors_from_tables():
    colors = set()
    for idx, table in enumerate(tables):
        for row in table.cells:
            for cell in row:
                colors.add(get_color(pdfs[idx],cell))
    return colors

colors =  {
    (255, 255, 0): 'GH', #gulur
    (198, 198, 198): '',#ljosgrar
    (240, 240, 240): '',#ljosljosgrar
    (255, 128, 255): 'NV',
    (0, 128, 0): 'BEG',#graenn
    (255, 0, 0): 'GR',#raudur
    (129, 129, 129): '',#mediumgrar
    (122, 122, 122): '',#mediumgrar
    (128, 0, 255): 'LRL',#fjolublar
    (255, 255, 255): 'UB',#hvitur
    (80, 138, 160): 'ORLOF',
    (128, 128, 64): '',
    (128, 128, 128): 'PHA',#mediumgrar
    (128, 0, 64): 'AS',
    }

def add_shift_text():
    for page_num,table in enumerate(tables):
        df = table.df
        for row_num,row in enumerate(table.cells):
            for col_num,cell in enumerate(row):
                cell_color = get_color(pdfs[page_num],cell=cell)
                if cell_color in colors and (cell.text != '') and re.match('[0-9][0-9]:[0-9][0-9]-[0-9][0-9]:[0-9][0-9]',cell.text):
                    df.iloc[row_num,col_num] += ' ' + colors[cell_color]
        # df.to_csv('/Users/odinndagur/Desktop/lolz' + str(page_num)+'.csv')

def old_main():
    parts = []
    for offset in range(0,len(tables),get_num_pages(tables)):
        first_page = tables[0+offset].df.iloc[1:,1:].copy()
        first_page.iloc[0,0] = 'Starfsmaður'
        second_page = tables[1+offset].df.copy()
        third_page = tables[2+offset].df.copy()

        first_page.index -= 1
        first_page.columns -= 1
        # second_page.index -= 1
        # third_page.index -= 1


        x,y = get_first_date_cell(second_page)
        first_page.iloc[0,1:] = second_page.iloc[y,x:]

        second_page = second_page.iloc[y+1:,x-1:]
        second_page.index -= 1 
        second_page.columns -= 1

        x,y = get_first_date_cell(third_page)
        third_page = third_page.iloc[y+1:,x-1:]
        third_page.index -= 1
        third_page.columns -= 1

        joined = pd.concat([first_page,second_page,third_page],ignore_index = True)
        joined.iloc[:,:]#.to_csv(os.path.join(output_directory,'out' + str(offset) + '.csv'))
        parts.append(joined)
    total = pd.concat(parts,axis=1,ignore_index=True)
    total.to_csv(os.path.join(output_directory,'output.csv'))

def process_df(table):
    df = table.df.copy()
    x,y = get_first_date_cell(df)
    df.iloc[y,x-1] = 'Starfsmaður'
    df.columns = df.iloc[y,:]
    df = df.iloc[y+1:,x-1:]
    df = df.set_index('Starfsmaður')
    df = df.replace('',nan).dropna(how='all').replace(nan,'')
    return df

add_shift_text()

processed_dfs = [process_df(table) for table in tables]

# concat fyrst
# concatenated_dfs = [str(offset) + str(idx) for offset in range(0,9,3) for idx in range(3)]
concatenated_dfs = [pd.concat(processed_dfs[offset:offset+get_num_pages(tables)]) for offset in range(0,tables.n,get_num_pages(tables))]

# síðan join
output_df = concatenated_dfs[0]
for df in concatenated_dfs[1:]:
    output_df = output_df.join(df)
print(f'saving to {output_directory} as {stripped_filename}.csv')
output_df.to_csv(os.path.join(output_directory,stripped_filename + '.csv'))

print(f'done saving, now making word documents')
for day in get_days(output_df):
    doc_from_date_day(day)
print(f'all done')

# pages = [pd.concat([*processed_dfs[offset:offset+get_num_pages(tables)]],axis=0,ignore_index=True) for offset in range(0,tables.n,get_num_pages(tables))]
# pd.concat([*pages],axis=1,ignore_index=True)
# parts = []
# def new_main():
#     for offset in range(0,len(tables),get_num_pages(tables)): #0, 3, 6
#         pgs = [process_df(table) for table in tables[offset:offset+get_num_pages(tables)]]
#         joined = pd.concat([*pgs],axis=0,ignore_index=True)
#         joined.reset_index()
#         parts.append(joined)
#     total = pd.concat([*parts],axis=1,ignore_index=True)
#     total.to_csv(os.path.join(output_directory,'output.csv'))
#     print(f'done, file output to {os.path.join(output_directory,"output.csv")}')




# new_main()


        # first_page = tables[0+offset].df.copy()
        # second_page = tables[1+offset].df.copy()
        # third_page = tables[2+offset].df.copy()

        # x,y = get_first_date_cell(first_page)
        # first_page.columns = first_page.iloc[y,:]
        # first_page = first_page.iloc[y+1:,x-1:]
        # first_page.reset_index(inplace=True, drop=True)

        # x,y = get_first_date_cell(second_page)
        # second_page.columns = second_page.iloc[y,:]
        # second_page = second_page.iloc[y+1:,x-1:]
        # second_page.reset_index(inplace=True, drop=True)

        # x,y = get_first_date_cell(third_page)
        # third_page.columns = third_page.iloc[y,:]
        # third_page = third_page.iloc[y+1:,x-1:]
        # # second_page = second_page.iloc[y+1:,x-1:]
        # third_page.reset_index(inplace=True, drop=True)


        # x,y = get_first_date_cell(third_page)
        # third_page = third_page.iloc[y+1:,x-1:]