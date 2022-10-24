#!/Users/odinndagur/.odb/vaktaplan/camelot-env/bin/python3


# IMPORTS
#region [blue2]
import camelot
import pandas as pd
import os
import sys
import re
from numpy import nan
from pathlib import Path
import pdfplumber
import datetime
from collections import defaultdict
import docx
from math import ceil,floor

import argparse
import sys
#endregion

# PARSE ARGS
#region [pink2]
parser = argparse.ArgumentParser()
parser.add_argument('file')
parser.add_argument('-p','--list-people', action='store_true', help='Print a list of people in the current shift plan')
parser.add_argument('-d','--dayplans', action='store_true', help='Should we generate docx dayplans for each day')
parser.add_argument('-o','--output-folder', help='Where should we save the dayplans (and csv if chosen)')
parser.add_argument('-n', '--name', help='Get shifts for specific person, add name of person')
parser.add_argument('-s','--save-csv', action='store_true', help='save intermediate table as csv')
args = parser.parse_args(sys.argv[1:])
#endregion


file = args.file
output_directory = args.output_folder if args.output_folder else os.path.join(os.path.dirname(file),'vaktaplan_output')
if not os.path.isdir(output_directory):
    os.makedirs(output_directory,exist_ok=True)
stripped_filename = os.path.splitext(os.path.basename(file))[0]
input_directory = file.split(os.path.basename(file))[0]

print(f'''
    running on file: {stripped_filename}
    from directory: {input_directory}
    outputting to: {output_directory}
''')

def main() -> None:
    if file.endswith('.pdf'):
        csv_path = os.path.join(output_directory,stripped_filename + '.csv')
        if os.path.exists(csv_path):
            print('csv exists in directory, using that')
            print('PSYCH not yet implemented')
        print('processing pdf')
        output_df = pdf_to_df(file)
        if args.save_csv:
            print(f'Saving to {output_directory} as {stripped_filename}.csv')
            output_df.to_csv(os.path.join(output_directory,stripped_filename + '.csv'))

    if file.endswith('.csv'):
        output_df = pd.read_csv(file,index_col=0,header=0)

    if args.dayplans:
        print(f'Generating docx files..')
        days = list(get_days(output_df))
        for idx, day in enumerate(days):
            temp_date = day['date']['date'].split('\n')[0]
            progress_bar_length = 100
            current_progress = ceil(idx/len(days) * progress_bar_length)
            progress_bar = f'[{"#" * current_progress}{" " * floor(progress_bar_length - current_progress)}]'
            print(f'\rMaking dayplan for {temp_date}...{progress_bar}',end='')
            doc_from_date_day(day)
        progress_bar = f'[{"#" * progress_bar_length}]'
        print(f'\rMaking dayplan for {temp_date}...{progress_bar}')
    if args.list_people:
        print('')
        for p in get_people(output_df):
            print(p)
        print('')
    
    if args.name:
        print()
        for shift in get_shifts_for_person(output_df,args.name):
            print(shift)
        print()
    os.system(f'open {output_directory}')
    print(f'All done!')



def pdf_to_df(file: str) -> pd.DataFrame:
    print('processing pdf')
    global h,w,new_h,new_w,pdfs
    from pdf2image import convert_from_path
    pdfs = convert_from_path(file)
    with pdfplumber.open(file) as pdf:
        page_1 = pdf.pages[0]
    h,w = page_1.height, page_1.width
    new_h,new_w = pdfs[0].height, pdfs[0].width

    tables = camelot.read_pdf(file,pages='1-end',flavor='lattice',line_scale=50,line_tol=1)
    add_shift_text(tables)
    processed_dfs = [process_df(table) for table in tables]
    # concat fyrst
    concatenated_dfs = [pd.concat(processed_dfs[offset:offset+get_num_pages(tables)]) for offset in range(0,tables.n,get_num_pages(tables))]
    # síðan join
    output_df = concatenated_dfs[0]
    for df in concatenated_dfs[1:]:
        output_df = output_df.join(df)
    return output_df

from ppl import ppl
def get_name(person: str):
    if person in ppl:
        if ppl[person]:
            return ppl[person]
    return person

def get_shifts_for_person(df: pd.DataFrame,person: str) -> str:
    shifts = df.loc[person].replace('',nan).dropna()
    return [f'{date.split(chr(10))[0]} {shift}' for date,shift in zip(shifts.index,shifts.values)]

def get_people(df: pd.DataFrame) -> list[str]:
    return [p for p in list(df.index) if len(p)]

def get_days(df: pd.DataFrame): #yields generator
    for idx in range(len(df.columns)):
        day = defaultdict(lambda: defaultdict(list))
        slice = df.iloc[:,idx].replace('',nan).dropna()
        if not len(slice):
            continue
        day['date']['date'] = slice.name
        # date = slice.name
        for row_idx, shift in enumerate(slice):
            person = get_name(slice.index[row_idx])
            if all([shift != 'ORLOF',shift]):
                time, type = shift.split(' ')
                if type == 'UB':
                    day['UB']['ub'].append(' '.join([person,time]))
                else:
                    starttime,endtime = time.split('-')
                    starthr = int(starttime.split(':')[0])
                    endhr = int(endtime.split(':')[0])
                    if 0 < starthr < 12:
                        if endhr > 16:
                            col = 'dv'
                        else:
                            col = 'mv'
                    if 12 <= starthr < 16:
                        col = 'dv'
                    if 16 <= starthr < 23:
                        col = 'kv'
                    if 23 <= starthr <= 24:
                        col = 'nv'
                    if 20 < endhr <= 24:
                        col = 'kv'
                    # print(person,starthr,endhr,type,col)
                    day[type][col].append(','.join([person,time,type,col]))
        yield day

def doc_from_date_day(day) -> docx.Document:
    date = day['date']['date']
    doc = docx.Document(os.path.join(os.path.dirname(os.path.abspath(__file__)),'fim_proto.docx'))
    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if cell.text.strip() == 'dags':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run(date.split('\n')[0])
            if cell.text.strip() == 'UB':
                p = cell.paragraphs[0]
                p.clear()
                p.add_run('\n'.join(day['UB']['ub']))
            if cell.text.strip() == 'vikudags':
                p = cell.paragraphs[0]
                p.clear()
                decimal_date = date.split('\n')[0]
                if decimal_date:
                    date_day, date_month = decimal_date.split('.')
                    weekday = get_weekday(month = date_month, day = date_day)
                    p.add_run(weekday)
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                shift = cell.text.split(' ')
                if not shift[0] in day:
                    cell.paragraphs[0].clear()
                else:
                    if shift[1] in day[shift[0]]:
                        temp = ' ' if 'NV' in shift[0] else '\n'
                        ppl = [
                            p.split(',')[0]+ temp + p.split(',')[1] for p in day[shift[0]][shift[1]]
                        ]
                        ppl = sorted(ppl, key=lambda x: int(x.split(temp)[-1].split(':')[0]))
                        p = cell.paragraphs[0]
                        p.clear()
                        p.add_run('\n'.join(ppl))


    for col_idx,col in enumerate(doc.tables[1].columns):
        for cell in doc.tables[1].column_cells(col_idx):
            if re.match('\D{2,3} [a-z]{2}',cell.text.strip()):
                cell.paragraphs[0].clear()
    date_str = date.split('\n')[0]
    output_filename = os.path.join(output_directory,f'{date_str} editad.docx')
    doc.save(output_filename)
    # os.system(f'open {os.path.join(os.getcwd(),output_filename)}')
    return doc

def get_weekday(year:int = None, month:int = None, day:int = None) -> str:
    days = ["Mánudagur", "Þriðjudagur", "Miðvikudagur", "Fimmtudagur", "Föstudagur", "Laugardagur", "Sunnudagur"]
    if not any([year,month,day]):
        raise ValueError('Need to input date')
    currentDate = datetime.datetime.today()
    if not year:
        if int(month) < currentDate.month and currentDate.month > 10:
            year = currentDate.year + 1
        else:
            year = currentDate.year
    inputDate = datetime.date(year=year,month=int(month),day=int(day))
    return days[inputDate.weekday()]

def is_first_page(df: pd.DataFrame) -> bool:
    for x in df.iloc[:,0].values:
        if 'Hæf' in x:
            return True
    return False

def get_first_date_cell(df: pd.DataFrame) -> tuple[int, int]:
    import re
    for row_idx,row in df.iterrows():
        for col_idx, cell in enumerate(row):
            # print(f'row {row_idx}, col {col_idx}, cell: {cell}')
            if re.match('[0-9][0-9]\.[0-9][0-9]', cell):
                return col_idx,row_idx

def get_num_pages(tables):
    counts = []
    last_first_page = 0
    for idx,table in enumerate(tables):
        if is_first_page(table.df):
            if idx > 0:
                counts.append(idx - last_first_page)
                last_first_page = idx
    if len(set(counts)) == 1:
        return counts[0]
    else:
        return counts

def get_color(img,cell):
    y = new_h - ((cell.y1 + 3)/h * new_h)
    x = ((cell.x1 + cell.x2)/2)/w * new_w
    return img.getpixel((x,y))

def get_colors_from_tables(tables):
    colors = set()
    for idx, table in enumerate(tables):
        for row in table.cells:
            for cell in row:
                colors.add(get_color(pdfs[idx],cell))
    return colors

colors =  {
    (255, 255, 0): 'GH', #gulur
    (198, 198, 198): '',#ljosgrar
    (240, 240, 240): '',#ljosljosgrar
    (255, 128, 255): 'NV',
    (0, 128, 0): 'BEG',#graenn
    (255, 0, 0): 'GR',#raudur
    (129, 129, 129): '',#mediumgrar
    (122, 122, 122): '',#mediumgrar
    (128, 0, 255): 'LRL',#fjolublar
    (255, 255, 255): 'UB',#hvitur
    (80, 138, 160): 'ORLOF',
    (128, 128, 64): '',
    (128, 128, 128): 'PHA',#mediumgrar
    (128, 0, 64): 'AS',
    }

def add_shift_text(tables) -> None:
    for page_num,table in enumerate(tables):
        df = table.df
        for row_num,row in enumerate(table.cells):
            for col_num,cell in enumerate(row):
                cell_color = get_color(pdfs[page_num],cell=cell)
                if cell_color in colors and (cell.text != '') and re.match('[0-9][0-9]:[0-9][0-9]-[0-9][0-9]:[0-9][0-9]',cell.text):
                    df.iloc[row_num,col_num] += ' ' + colors[cell_color]

def process_df(table) -> pd.DataFrame:
    df = table.df.copy()
    x,y = get_first_date_cell(df)
    df.iloc[y,x-1] = 'Starfsmaður'
    df.columns = df.iloc[y,:]
    df = df.iloc[y+1:,x-1:]
    df = df.set_index('Starfsmaður')
    df = df.replace('',nan).dropna(how='all').dropna(how='all',axis=1).replace(nan,'')
    return df

if __name__ == '__main__':
    main()